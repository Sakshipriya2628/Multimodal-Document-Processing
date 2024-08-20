import os
import base64
import logging
from docx import Document
from dotenv import load_dotenv, find_dotenv
from openai import AzureOpenAI

# Load environment variables
load_dotenv()

# Retrieve API keys and endpoint URL from environment variables
api_key = os.environ.get('AZURE_VISION_API_KEY')
api_url = os.environ.get('vision_base_url')

# Configure AzureOpenAI client
client = AzureOpenAI(
    azure_endpoint=api_url, 
    api_key=api_key,  
    api_version="2023-12-01-preview"
)
model = "gpt4-vision-preview"

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# Function to get file paths from a folder
def get_file_paths(folder_path):
    file_paths = []
    for filename in sorted(os.listdir(folder_path)):
        file_path = os.path.join(folder_path, filename)
        if os.path.isfile(file_path):
            file_paths.append(file_path)
    return file_paths

# Function to prepare image data for analysis
def image_selection(file_path):
    with open(file_path, 'rb') as image_file:
        binary_content = image_file.read()
    base64_encoded_str = base64.b64encode(binary_content).decode('utf-8')
    data_url = f'data:image/png;base64,{base64_encoded_str}'
    messages = [
        {
            "role": "user",
            "content": [
                {"type": "text", "text": "Detailed image analysis instructions..."},
                {"type": "image_url", "image_url": {"url": data_url}}
            ]
        }
    ]
    return messages

# Function to describe images and output to a Word document
def describe_images(folder_path):
    file_paths = get_file_paths(folder_path)
    data = []
    doc = Document()  # Create a new Word document

    for file_path in file_paths:
        try:
            logging.info(f"Starting analysis for Image {file_paths.index(file_path) + 1}")
            messages = image_selection(file_path)
            resp = client.chat.completions.create(
                model=model,
                messages=messages,
                max_tokens=4000,
                seed=42,
            )
            description = resp.choices[0].message.content
            data.append({
                'Image Number': file_paths.index(file_path) + 1,
                'Image Path': file_path,
                'Description': description
            })
            logging.info("End of analysis")
        except Exception as e:
            logging.error(f"Failed to process file {file_path}: {str(e)}")

    if data:
        for entry in data:
            doc.add_paragraph(f"Image Number: {entry['Image Number']}")
            doc.add_paragraph(f"Description: {entry['Description']}")
            doc.add_paragraph("--------------------------------------------------")
        doc.save('Baltic 5 Analysis.docx')  # Save the document
        logging.info("Data successfully written to Word document.")
    else:
        logging.warning("No data collected. Check the input folder and image processing.")

# Main function to execute
if __name__ == "__main__":
    folder_path = '/Users/sakshi_admin/Desktop/Image Assessment POC/Baltic 5 output_images'
    describe_images(folder_path)








